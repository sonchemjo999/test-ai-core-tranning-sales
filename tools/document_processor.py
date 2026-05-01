"""
Document Processor — Phase 1 of Hybrid NotebookLM Architecture.

Converts enterprise documents (PDF, DOCX, XLSX/CSV) into clean Markdown text,
preserving table structures and headings for downstream LLM consumption.
"""

from __future__ import annotations

import csv
import io
import logging
import uuid
from pathlib import Path
from typing import Literal

logger = logging.getLogger(__name__)

# Supported MIME types → internal format keys
SUPPORTED_MIMES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "text/csv": "csv",
    "text/plain": "txt",
}

SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".csv": "csv",
    ".txt": "txt",
}


class DocumentProcessingError(Exception):
    """Raised when document parsing fails."""
    pass


def detect_format(filename: str, content_type: str | None = None) -> str:
    """Detect document format from filename extension or MIME type."""
    ext = Path(filename).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[ext]
    if content_type and content_type in SUPPORTED_MIMES:
        return SUPPORTED_MIMES[content_type]
    raise DocumentProcessingError(
        f"Unsupported file format: {filename} (content_type={content_type}). "
        f"Supported: PDF, DOCX, XLSX, CSV, TXT."
    )


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF, preserving layout structure."""
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.layout import LAParams

    laparams = LAParams(
        line_margin=0.5,
        word_margin=0.1,
        char_margin=2.0,
        boxes_flow=0.5,
    )
    try:
        text = pdfminer_extract(io.BytesIO(file_bytes), laparams=laparams)
    except Exception as e:
        raise DocumentProcessingError(f"PDF parsing failed: {e}") from e

    if not text or not text.strip():
        raise DocumentProcessingError(
            "Could not extract text from this PDF. "
            "It may be image-only (scanned). Please provide a text-based PDF."
        )
    return text.strip()


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX, converting headings to Markdown."""
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise DocumentProcessingError(f"DOCX parsing failed: {e}") from e

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "heading" in style_name:
            lines.append(f"#### {text}")
        elif "list" in style_name or "bullet" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Also extract tables from DOCX
    for table in doc.tables:
        lines.append("")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    result = "\n".join(lines).strip()
    if not result:
        raise DocumentProcessingError("DOCX file appears to be empty.")
    return result


def _parse_xlsx(file_bytes: bytes) -> str:
    """Convert XLSX sheets into Markdown tables."""
    from openpyxl import load_workbook

    try:
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as e:
        raise DocumentProcessingError(f"XLSX parsing failed: {e}") from e

    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        section_lines = [f"## {sheet_name}"]

        # Filter out completely empty rows
        non_empty_rows = [r for r in rows if any(c is not None for c in r)]
        if not non_empty_rows:
            continue

        # First non-empty row is the header
        header = non_empty_rows[0]
        max_cols = len(header)
        header_strs = [str(c or "").strip().replace("|", "\\|") for c in header]
        section_lines.append("| " + " | ".join(header_strs) + " |")
        section_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        for row in non_empty_rows[1:]:
            cells = [str(c or "").strip().replace("|", "\\|") for c in row[:max_cols]]
            # Pad if row is shorter than header
            while len(cells) < max_cols:
                cells.append("")
            section_lines.append("| " + " | ".join(cells) + " |")

        sections.append("\n".join(section_lines))

    wb.close()

    result = "\n\n".join(sections).strip()
    if not result:
        raise DocumentProcessingError("XLSX file appears to be empty.")
    return result


def _parse_csv(file_bytes: bytes) -> str:
    """Convert CSV into a Markdown table."""
    try:
        text = file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    reader = csv.reader(io.StringIO(text))
    rows = [r for r in reader if any(c.strip() for c in r)]

    if not rows:
        raise DocumentProcessingError("CSV file appears to be empty.")

    lines: list[str] = []
    header = rows[0]
    max_cols = len(header)
    header_strs = [c.strip().replace("|", "\\|") for c in header]
    lines.append("| " + " | ".join(header_strs) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    for row in rows[1:]:
        cells = [c.strip().replace("|", "\\|") for c in row[:max_cols]]
        while len(cells) < max_cols:
            cells.append("")
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def _parse_txt(file_bytes: bytes) -> str:
    """Read plain text as-is."""
    try:
        text = file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    if not text.strip():
        raise DocumentProcessingError("Text file appears to be empty.")
    return text.strip()


# Parser dispatch table
_PARSERS = {
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "xlsx": _parse_xlsx,
    "csv": _parse_csv,
    "txt": _parse_txt,
}


def _estimate_token_count(text: str) -> int:
    """Rough token count estimate (~0.75 words per token for Vietnamese/mixed text)."""
    word_count = len(text.split())
    return int(word_count / 0.75)


def process_document(
    file_bytes: bytes,
    filename: str,
    content_type: str | None = None,
) -> dict:
    """
    Main entry point — Process a raw document into clean Markdown.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used for format detection).
        content_type: Optional MIME type from the upload header.

    Returns:
        dict with keys:
            - document_id: str (UUID)
            - filename: str
            - format: str (pdf/docx/xlsx/csv/txt)
            - raw_markdown: str (Clean Markdown text)
            - token_count_estimate: int
    """
    fmt = detect_format(filename, content_type)
    logger.info(f"Processing document: {filename} (format={fmt}, size={len(file_bytes)} bytes)")

    parser = _PARSERS[fmt]
    raw_markdown = parser(file_bytes)

    # Basic cleanup: collapse excessive whitespace, normalize line breaks
    lines = raw_markdown.splitlines()
    cleaned_lines: list[str] = []
    prev_empty = False
    for line in lines:
        is_empty = not line.strip()
        if is_empty and prev_empty:
            continue  # Skip consecutive empty lines
        cleaned_lines.append(line)
        prev_empty = is_empty

    raw_markdown = "\n".join(cleaned_lines).strip()
    token_estimate = _estimate_token_count(raw_markdown)

    logger.info(
        f"Document processed: {filename} → {len(raw_markdown)} chars, "
        f"~{token_estimate} tokens"
    )

    return {
        "document_id": str(uuid.uuid4()),
        "filename": filename,
        "format": fmt,
        "raw_markdown": raw_markdown,
        "token_count_estimate": token_estimate,
    }
